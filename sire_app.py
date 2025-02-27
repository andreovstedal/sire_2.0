import streamlit as st
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches
import io
import tempfile

st.set_page_config(
    page_title="SIRE 2.0 Report",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for styling
st.markdown("""
<style>
    /* Base styles */
    .main {
        background-color: #f0f0f0;
        padding: 1rem 2rem;
    }
    .stApp {
        background-color: #f0f0f0;
    }
    
    /* Remove sidebar */
    [data-testid="stSidebar"] {
        background-color: #f0f0f0;
        width: 0px !important;
        min-width: 0px !important;
        flex: 0 !important;
        -webkit-box-flex: 0 !important;
    }
    
    /* Header styling */
    .header {
        background-color: #5c7cba;
        padding: 20px;
        border-radius: 5px;
        color: white;
        margin-bottom: 20px;
    }
    
    /* Section headers */
    h2 {
        color: #5c7cba;
        font-weight: bold;
        margin-top: 20px;
        margin-bottom: 10px;
    }
    
    /* Container styling */
    [data-testid="stVerticalBlock"] {
        background-color: white;
        border-radius: 5px;
        padding: 1rem;
        margin-bottom: 1rem;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    
    /* Button styling */
    .stButton > button {
        background-color: #5c7cba;
        color: white;
        font-weight: 600;
        border: none;
    }
    .stButton > button:hover {
        background-color: #4a69a7;
    }
    
    /* Footer copyright */
    .footer {
        position: fixed;
        bottom: 0;
        left: 0;
        right: 0;
        background-color: #5c7cba;
        color: rgba(255, 255, 255, 0.7);
        text-align: center;
        padding: 5px;
        font-size: 12px;
        z-index: 999;
    }
</style>

<!-- Header -->
<div class="header">
    <h1>SIRE 2.0 Report</h1>
</div>
""", unsafe_allow_html=True)

# Utility functions
def format_date(date_str: str) -> str:
    """Format dates to the desired format"""
    if not date_str:
        return ""
    
    try:
        if "T" in date_str:
            dt = datetime.strptime(date_str.split(".")[0], "%Y-%m-%dT%H:%M:%S")
            return dt.strftime("%Y-%m-%d %H:%M")
        return date_str
    except Exception:
        return date_str

def generate_question_numbers(comments: list) -> dict:
    """Generate sequential numbers for question IDs"""
    question_map = {}
    current_group = 0
    current_number = 0
    last_id = None
    
    for comment in comments:
        current_id = comment['id']
        
        if current_id not in question_map:
            if current_id != last_id:
                current_group += 1
                current_number = 1
            question_map[current_id] = f"{current_group}.{current_number}"
            last_id = current_id
        current_number += 1
    
    return question_map

def process_inspection_data(inspection_data):
    """Process the loaded inspection data"""
    if not inspection_data:
        return None, None
        
    try:
        # Process metadata
        metadata = {}
        for item in inspection_data.get('metaData', []):
            key = item['key']
            value = item['value']
            
            # Format dates in metadata
            if any(x in key for x in ['DATE', 'DATETIME']):
                value = format_date(value)
            
            metadata[key] = value

        metadata_list = list(metadata.items())

        # Process comments
        comments = []
        for question in inspection_data.get('questions', []):
            template_id = question.get('templateQuestionId', '')
            
            for response in question.get('complexResponses', []):
                for observation in response.get('observations', []):
                    if observation.get('comments'):
                        for op_comment in observation.get('initialOperatorComments', []):
                            comments.append({
                                'id': template_id,
                                'inspector_comment': observation.get('comments', ''),
                                'operator_comment': op_comment.get('comments', ''),
                                'date': format_date(op_comment.get('commentDate', ''))
                            })

        # Generate sequential question numbers
        question_numbers = generate_question_numbers(comments)
        
        # Create comments list
        comments_data = []
        for comment in comments:
            # Use the mapped question number instead of the UUID
            question_number = question_numbers[comment['id']]
            comments_data.append([
                question_number,
                comment['inspector_comment'],
                comment['operator_comment'],
                comment['date']
            ])
        
        return metadata_list, comments_data
        
    except Exception as e:
        st.error(f"Error processing data: {str(e)}")
        return None, None

def create_docx(metadata_list, comments_data):
    """Create a Word document with the inspection data"""
    doc = Document()
    
    # Add title
    doc.add_heading('SIRE 2.0 Inspection Report', 0)
    
    # Add metadata section
    doc.add_heading('Vessel Information', level=1)
    metadata_table = doc.add_table(rows=1, cols=2)
    metadata_table.style = 'Table Grid'
    
    header_cells = metadata_table.rows[0].cells
    header_cells[0].text = 'Field'
    header_cells[1].text = 'Value'
    
    for key, value in metadata_list:
        row_cells = metadata_table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)
    
    # Add comments section
    doc.add_heading('Comments and Observations', level=1)
    comments_table = doc.add_table(rows=1, cols=4)
    comments_table.style = 'Table Grid'
    
    header_cells = comments_table.rows[0].cells
    headers = ["Question", "Inspector Comment", "Operator Comment", "Date"]
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    # Set column widths (proportional)
    widths = [1, 3, 3, 1.5]  # Relative widths
    for i, width in enumerate(widths):
        for cell in comments_table.columns[i].cells:
            cell.width = Inches(width)
    
    for row_data in comments_data:
        row_cells = comments_table.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = str(cell_value)
    
    # Save to BytesIO object
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

# Initialize session state
if 'inspection_data' not in st.session_state:
    st.session_state.inspection_data = None
if 'metadata_list' not in st.session_state:
    st.session_state.metadata_list = None
if 'comments_data' not in st.session_state:
    st.session_state.comments_data = None

# Main interface
col1, col2 = st.columns([1, 4])

with col1:
    uploaded_file = st.file_uploader("Upload JSON File", type=['json'])
    
    if st.button("Process File"):
        if uploaded_file is not None:
            try:
                # Read and clean the file content
                file_content = uploaded_file.read().decode('utf-8', errors='replace')
                
                # Clean up the JSON content
                # Remove extra quotes at the start
                if file_content.startswith('{",'):
                    file_content = file_content.replace('{",', '{', 1)
                elif file_content.startswith('{"'):
                    pass  # This is correct JSON start
                elif file_content.startswith('{,'):
                    file_content = file_content.replace('{,', '{', 1)
                
                # Remove any BOM or whitespace
                file_content = file_content.strip()
                
                # Additional cleanup
                file_content = file_content.replace('\x00', '')  # Remove null bytes
                file_content = file_content.replace('\r', '')    # Remove carriage returns
                
                try:
                    inspection_data = json.loads(file_content)
                    st.session_state.inspection_data = inspection_data
                    st.session_state.metadata_list, st.session_state.comments_data = process_inspection_data(inspection_data)
                    st.success("File loaded successfully!")
                except json.JSONDecodeError as e:
                    # Try more aggressive cleaning
                    try:
                        # Remove any non-standard characters
                        clean_content = ''.join(c for c in file_content if ord(c) >= 32 or c in '\n\t')
                        # Ensure proper JSON structure
                        if not clean_content.startswith('{'):
                            clean_content = '{' + clean_content.split('{', 1)[1]
                        inspection_data = json.loads(clean_content)
                        st.session_state.inspection_data = inspection_data
                        st.session_state.metadata_list, st.session_state.comments_data = process_inspection_data(inspection_data)
                        st.success("File loaded successfully with cleanup!")
                    except json.JSONDecodeError as e:
                        st.error(f"Invalid JSON format: {str(e)}")
            except Exception as e:
                st.error(f"Error loading file: {str(e)}")
        else:
            st.warning("Please upload a JSON file first")

    if st.session_state.metadata_list is not None and st.session_state.comments_data is not None:
        # Export button
        docx_bytes = create_docx(st.session_state.metadata_list, st.session_state.comments_data)
        st.download_button(
            label="Export to DOCX",
            data=docx_bytes,
            file_name="SIRE_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

with col2:
    if st.session_state.metadata_list is not None:
        st.subheader("Vessel Information")
        
        # Create a custom table for metadata without pandas
        metadata_html = "<table width='100%' style='border-collapse: collapse;'>"
        metadata_html += "<tr><th style='border: 1px solid #ddd; padding: 8px; text-align: left; background-color: #5c7cba; color: white;'>Field</th>"
        metadata_html += "<th style='border: 1px solid #ddd; padding: 8px; text-align: left; background-color: #5c7cba; color: white;'>Value</th></tr>"
        
        for key, value in st.session_state.metadata_list:
            metadata_html += f"<tr><td style='border: 1px solid #ddd; padding: 8px;'>{key}</td>"
            metadata_html += f"<td style='border: 1px solid #ddd; padding: 8px;'>{value}</td></tr>"
        
        metadata_html += "</table>"
        st.markdown(metadata_html, unsafe_allow_html=True)
    
    if st.session_state.comments_data is not None:
        st.subheader("Comments and Observations")
        
        # Create a custom table for comments without pandas
        comments_html = "<table width='100%' style='border-collapse: collapse;'>"
        comments_html += "<tr><th style='border: 1px solid #ddd; padding: 8px; text-align: left; background-color: #5c7cba; color: white;'>Question</th>"
        comments_html += "<th style='border: 1px solid #ddd; padding: 8px; text-align: left; background-color: #5c7cba; color: white;'>Inspector Comment</th>"
        comments_html += "<th style='border: 1px solid #ddd; padding: 8px; text-align: left; background-color: #5c7cba; color: white;'>Operator Comment</th>"
        comments_html += "<th style='border: 1px solid #ddd; padding: 8px; text-align: left; background-color: #5c7cba; color: white;'>Date</th></tr>"
        
        for row in st.session_state.comments_data:
            comments_html += f"<tr><td style='border: 1px solid #ddd; padding: 8px;'>{row[0]}</td>"
            comments_html += f"<td style='border: 1px solid #ddd; padding: 8px;'>{row[1]}</td>"
            comments_html += f"<td style='border: 1px solid #ddd; padding: 8px;'>{row[2]}</td>"
            comments_html += f"<td style='border: 1px solid #ddd; padding: 8px;'>{row[3]}</td></tr>"
        
        comments_html += "</table>"
        st.markdown(comments_html, unsafe_allow_html=True)

# Footer
st.markdown("""
<div class="footer">
    SIRE 2.0 Report Tool
</div>
""", unsafe_allow_html=True)
